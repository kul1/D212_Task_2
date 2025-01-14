from docx import Document  # Import the Document class
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from datetime import date

# Define hardcoded configuration parameters
CONFIG_TYPE = "pca_analysis"
VISUALS_DIR = "results/pca_analysis/visuals"
OUTPUT_PATH = "results/pca_analysis/PCA_Report.docx"

print(f"Loaded configuration for {CONFIG_TYPE}.")

# Utility functions
def add_code_snippet(document, code_text):
    """Add a formatted code snippet to the document."""
    p = document.add_paragraph()
    r = p.add_run(code_text)
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_visual(document, image_path, caption):
    """Add a visual with caption."""
    if os.path.exists(image_path):
        document.add_picture(image_path, width=Inches(5.5))
        p = document.add_paragraph(caption)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        p = document.add_paragraph(f"Visual not found: {caption}")
        p.italic = True

def add_apa_references(document, references):
    """Add APA-style references."""
    document.add_heading("References", level=1)
    for ref in references:
        citation = f"{ref['author']} ({ref['year']}). {ref['title']}. Retrieved from {ref['source']}"
        document.add_paragraph(citation, style="Normal")

def create_report():
    document = Document()

    # Title Page
    document.add_heading("Data Mining II — D212", level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph("Task 2: Principal Component Analysis (PCA)", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph("by").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph("Prateep Kul").bold = True
    document.add_paragraph(f"{date.today().strftime('%B %d, %Y')}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_page_break()

    # Table of Contents
    document.add_heading("Table of Contents", level=1)
    toc = [
        "Part I: Research Question",
        "Part II: Method Justification",
        "Part III: Data Preparation",
        "Part IV: Analysis",
        "Webguide",
        "References"
    ]
    for item in toc:
        document.add_paragraph(item, style="List Number")
    document.add_page_break()

    # Part I
    document.add_heading("Part I: Research Question", level=1)
    document.add_paragraph(
        "The research question focuses on understanding the factors contributing to overweight patients, "
        "significantly impacting cost-effective treatment plans through PCA. This analysis utilizes a dataset "
        "containing demographic, medical, and hospital service information from 10,000 patients. PCA will help "
        "identify key patterns and correlations among the features. (Pedregosa et al., 2011)."
    )
    document.add_page_break()

    # Part II
    document.add_heading("Part II: Method Justification", level=1)
    document.add_paragraph(
        "Principal Component Analysis (PCA) was selected for this analysis due to its ability to reduce dimensionality "
        "while retaining essential information. By transforming features into uncorrelated principal components, PCA "
        "improves interpretability, mitigates multicollinearity, and enhances computational efficiency (Waskom, 2017)."
    )
    document.add_page_break()

    # Part III: Data Preparation
    document.add_heading("Part III: Data Preparation", level=1)

    # Section 1: Continuous Variables for PCA Analysis
    document.add_heading("1. Continuous Variables for PCA Analysis", level=2)
    document.add_paragraph(
        "To address the research question regarding factors contributing to overweight patients, "
        "the following continuous variables were selected for PCA analysis. These variables were chosen "
        "to capture relevant aspects of patient demographics, dietary habits, and medical history:"
    )
    document.add_paragraph(
        "1. Age (Continuous): Represents the patient's age in years.\n"
        "2. Income (Continuous): Annual income of the patient in USD.\n"
        "3. Full Meals Eaten per Day (Continuous): The average number of full meals consumed daily.\n"
        "4. Soft Drink Consumption (Continuous): The average number of soft drinks consumed per day.\n"
        "5. Vitamin D Levels (VitD_levels) (Continuous): The patient's blood vitamin D concentration in ng/mL.\n"
        "6. Doctor Visits (Doc_visits) (Continuous): Number of medical visits in the past year."
    )
    document.add_paragraph(
        "These variables were confirmed as continuous measures through detailed inspection of the dataset and the configuration code. "
        "Their inclusion aligns with the goal of understanding patterns related to overweight status, a critical factor in cost-effective treatment planning."
    )

    # Section 2: Steps for Data Preparation
    document.add_heading("2. Steps for Data Preparation", level=2)
    document.add_paragraph(
        "To ensure the data was ready for PCA analysis, the following data cleaning and preparation steps were implemented:"
    )
    document.add_paragraph(
        "1. General Cleaning Steps:\n"
        "- Handling Missing Values: Missing values in the dataset were imputed using the mean for continuous variables.\n"
        "- Outlier Treatment: Outliers were capped within 1.5 times the interquartile range (IQR) to reduce their impact on PCA results.\n"
        "- Data Type Verification: The data types for all selected variables were verified to confirm they were continuous measures.\n"
    )
    document.add_paragraph(
        "2. Specific Preparation for PCA Analysis:\n"
        "- Standardization: Continuous variables were standardized to have a mean of 0 and a standard deviation of 1. This step ensures "
        "that all variables contribute equally to the PCA model, avoiding bias caused by scale differences.\n"
        "- Validation of Continuous Variables: Only the variables explicitly needed for the PCA analysis of overweight status were retained. "
        "This careful selection minimizes noise and focuses the analysis on relevant patterns."
    )

    # Section 3: Visualization of Standardization
    document.add_heading("3. Visualization of Standardization", level=2)
    document.add_paragraph(
        "The figures below demonstrate the distribution of features before and after standardization, "
        "emphasizing the importance of this preprocessing step:"
    )
    add_visual(document, f"{VISUALS_DIR}/original_distribution.png", "Figure 2: Distribution of Features (Before Standardization)")
    add_visual(document, f"{VISUALS_DIR}/standardized_distribution.png", "Figure 3: Distribution of Features (After Standardization)")
    document.add_paragraph(
        "Figure 2 shows that the raw data for continuous variables displays varying scales and distributions, which can bias PCA results. "
        "For example, Income spans a much broader range than Age, leading to disproportionate influence in the analysis.\n"
        "Figure 3 shows that after standardization, all variables exhibit comparable distributions centered around a mean of 0 with a standard deviation of 1. "
        "This transformation ensures that PCA focuses on variance patterns rather than differences in scale, aligning with its mathematical foundation."
    )
    # Section 4: Explained Variance by PCA Components
    document.add_heading("4. Explained Variance by PCA Components", level=2)
    document.add_paragraph(
        "Figure 4 illustrates the explained variance by each principal component, along with the cumulative variance across all components."
    )
    document.add_paragraph(
        "- Individual Variance: The bar chart represents the proportion of variance captured by each individual principal component. "
        "For instance, the first principal component (PC1) captures 8.98% of the total variance, the highest among all components.\n"
        "- Cumulative Variance: The line graph shows how variance accumulates as additional components are included. "
        "The cumulative variance reaches 33.37% when the first five principal components are considered."
    )
    document.add_paragraph(
        "This analysis is critical because:\n"
        "1. Feature Reduction: The figure shows that only a small subset of components (e.g., the first five) captures a significant portion of the variance, "
        "allowing dimensionality reduction without substantial information loss.\n"
        "2. Focus on Key Patterns: By retaining the components that explain the majority of the variance, we can focus on the most meaningful patterns in the data, "
        "simplifying the model and improving interpretability.\n"
        "3. Overweight Analysis Context: In the context of understanding overweight status, retaining these five components ensures that the underlying factors contributing "
        "to variance are preserved while discarding noise."
    )
    add_visual(document, f"{VISUALS_DIR}/explained_variance.png", "Figure 4: Explained Variance by PCA Components")


    # Part IV: Analysis
    document.add_heading("Part IV: Analysis", level=1)
    document.add_paragraph(
        "Principal Components Analysis (PCA) was performed to identify critical components that explain the variance "
        "in the dataset, focusing on overweight patients and their related variables."
    )

    # Matrix of Principal Components
    document.add_heading("1. Matrix of Principal Components", level=2)
    document.add_paragraph(
        "The principal component matrix highlights the relationships between original variables and their contributions "
        "to the PCA components. Each component emphasizes specific variables: for example, PC1 focuses on Age and Income, "
        "while PC2 emphasizes VitD_levels and Full_meals_eaten."
    )
    add_visual(document, f"{VISUALS_DIR}/correlation_heatmap.png", "Figure 1: Correlation Heatmap")

    # Number of Principal Components Retained
    document.add_heading("2. Number of Principal Components Retained", level=2)
    document.add_paragraph(
        "Using the elbow criterion, five principal components were retained as they captured significant variance  (Pedregosa et al., 2011)."
        "without overfitting the data."
    )
    add_visual(document, f"{VISUALS_DIR}/scree_plot_with_elbow.png", "Figure 2: Scree Plot with Elbow Point")

    # Explained Variance
    document.add_heading("3. Explained Variance by Components", level=2)
    document.add_paragraph(
        "The table below shows the variance explained by each retained component. Together, these components capture "
        "33.37% of the total variance."
    )
    add_code_snippet(document, """
PC1: 8.98% (Cumulative: 8.98%)
PC2: 8.37% (Cumulative: 17.35%)
PC3: 5.78% (Cumulative: 23.13%)
PC4: 5.22% (Cumulative: 28.36%)
PC5: 5.01% (Cumulative: 33.37%)
""")

    # Total Variance Captured
    document.add_heading("4. Total Variance Captured", level=2)
    document.add_paragraph(
        "The five retained components collectively account for 33.37% of the variance, effectively reducing dimensionality "
        "while preserving key information about overweight patients."
    )
    add_visual(document, f"{VISUALS_DIR}/explained_variance.png", "Figure 3: Explained Variance and Cumulative Variance")

    # Summary of Results
    document.add_heading("5. Summary of Results", level=2)
    document.add_paragraph(
        "The PCA analysis successfully reduced the dataset's dimensionality to five principal components. Key findings include:\n"
        "- PC1: Strongly associated with Age and Income.\n"
        "- PC2: Highlights the importance of VitD_levels and Full_meals_eaten.\n"
        "- PC3–PC5: Capture additional variance, including Doc_visits and TotalCharge.\n\n"
        "These findings provide insights into the factors contributing to overweight patients, guiding cost-effective treatment plans."
    )

    # Supporting Visualizations
    document.add_heading("Supporting Visualizations", level=1)
    add_visual(document, f"{VISUALS_DIR}/Age_distribution.png", "Figure 4: Distribution of Age")
    add_visual(document, f"{VISUALS_DIR}/Income_distribution.png", "Figure 5: Distribution of Income")
    add_visual(document, f"{VISUALS_DIR}/VitD_levels_distribution.png", "Figure 6: Distribution of VitD_levels")
    add_visual(document, f"{VISUALS_DIR}/Doc_visits_distribution.png", "Figure 7: Distribution of Doc_visits")
    add_visual(document, f"{VISUALS_DIR}/standardized_distribution.png", "Figure 8: Standardized Data Distribution")

    document.add_page_break()

    # Webguide and References
    document.add_heading("Webguide", level=1)
    document.add_paragraph(
        "The following resources provide a deeper understanding of PCA and its applications:"
    )
    document.add_paragraph("- https://scikit-learn.org/stable/modules/generated/sklearn.decomposition.PCA.html")
    document.add_paragraph("- https://towardsdatascience.com/a-guide-to-principal-component-analysis-8727221e5d96")
    document.add_paragraph("- https://machinelearningmastery.com/principal-components-analysis-for-dimensionality-reduction/")
    document.add_page_break()

    add_apa_references(document, [
        {"author": "Pedregosa, F. et al.", "year": "2011", "title": "Scikit-learn: Machine Learning in Python",
         "source": "https://jmlr.org/papers/v12/pedregosa11a.html"},
        {"author": "Waskom, M. L.", "year": "2017", "title": "Seaborn: Statistical Data Visualization",
         "source": "https://seaborn.pydata.org/"},
    ])

    # Save the document
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"Report saved to: {OUTPUT_PATH}")

if __name__ == "__main__":
    create_report()
